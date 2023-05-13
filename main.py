
import datetime
from os import name
import os
import sys
import platform
from PyQt5.QtCore import QTimer, pyqtSlot
from PyQt5.QtWidgets import QApplication, QFileDialog, QLabel, QMainWindow
from PyQt5.QtGui import QPixmap, QImage
from PySide2 import QtCore, QtGui, QtWidgets
from PySide2.QtCore import (QCoreApplication, QPropertyAnimation, QDate, QDateTime, QMetaObject, QObject, QPoint, QRect, QSize, QTime, QUrl, Qt, QEvent)
from PySide2.QtGui import (QBrush, QColor, QConicalGradient, QCursor, QFont, QFontDatabase, QIcon, QKeySequence, QLinearGradient, QPalette, QPainter, QPixmap, QRadialGradient)
from PySide2.QtWidgets import *
import pandas as pd
from PyQt5.QtGui import QPixmap

from os.path import exists
import time
import csv
import docx
from fpdf import FPDF
from PyQt5.QtWidgets import QMessageBox
# GUI FILE
from ui_main import Ui_MainWindow
from datetime import datetime, timedelta
import re

# pyinstaller --onefile --windowed --icon=logo.ico .\main.py
from main_function import *


class MainWindow(QMainWindow):
    def __init__(self):
        QMainWindow.__init__(self)
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.ui.btnUpload.clicked.connect(self.getfiles)
        self.ui.btnProcess.clicked.connect(self.processAudio)
        self.ui.btnClear.clicked.connect(self.clearText)

        self.ui.btnDocs.clicked.connect(self.exportToDocs)
        self.ui.btnPDF.clicked.connect(self.exportPDF)

        self.ui.btnStartLive.clicked.connect(self.startLive)
        self.ui.btnClearLive.clicked.connect(self.clearTextLive) 
        self.ui.btnStopLive.clicked.connect(self.stopRecording) 

        self.ui.btnDocsLive.clicked.connect(self.exportToDocsLive)
        self.ui.btnPDFLive.clicked.connect(self.exportPDFLive)
        self.ui.btnSrt.clicked.connect(self.exportToSRT)

        self.show()

            ########################################################################

        # PAGE 1
        self.ui.btn_page_1.clicked.connect(lambda: self.ui.stackedWidget.setCurrentWidget(self.ui.page_1))

        # PAGE 2
        self.ui.btn_page_2.clicked.connect(lambda: self.ui.stackedWidget.setCurrentWidget(self.ui.page_2))

        # PAGE 3
        self.ui.btn_page_3.clicked.connect(lambda: self.ui.stackedWidget.setCurrentWidget(self.ui.page_3))

        # PAGE 4
        p = pyaudio.PyAudio()
        # Get list of available audio devices
        info = p.get_host_api_info_by_index(0)
        numdevices = info.get('deviceCount')
        for i in range(0, numdevices):
                if (p.get_device_info_by_host_api_device_index(0, i).get('maxInputChannels')) > 0:
                    print("Input Device id ", i, " - ", p.get_device_info_by_host_api_device_index(0, i).get('name'))
                    self.ui.selectMicrophone.addItem(p.get_device_info_by_host_api_device_index(0, i).get('name'))
        
        self.ui.btnStopLive.setEnabled(False)
        self.ui.btnStopLive.setStyleSheet("background-color: #834545;")


                
        main_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
        ffmpeg_bin_dir = os.path.join(main_dir, 'ffmpeg', 'bin')
        
        if ffmpeg_bin_dir not in os.environ['PATH']:
            os.environ['PATH'] += os.pathsep + ffmpeg_bin_dir



    def getfiles(self):
        fileName, _ = QtWidgets.QFileDialog.getOpenFileName(self, 'Single File', QtCore.QDir.rootPath() , '*')
        self.ui.uploadFilename.setText(fileName)
        self.filePath = fileName

    def processAudio(self):
        self.ui.outputText.setText(' ')
        filename = self.ui.uploadFilename.text()
        # destination_folder = os.path.dirname(os.path.abspath(__file__))
        # if not os.path.exists(destination_folder):
        #     os.makedirs(destination_folder)
        # os.replace(self.filePath, destination_folder+ "/uploads/" + QtCore.QFileInfo(self.filePath).fileName())
        if not filename:
                # Display an error message
            QtWidgets.QMessageBox.critical(self, "Error", "No file selected.")
            return
  
        self.ui.consoleLog.append("> Uploading Audio") 
        time.sleep(3)
        self.ui.consoleLog.append("> Upload Complete")
        time.sleep(3) 
        self.ui.consoleLog.append("> Initializing Model") 
        time.sleep(3)
        # Show the loading screen

        if self.ui.enableSpeakerDiarization.isChecked():
            msgBox = QMessageBox()
            msgBox.setIcon(QMessageBox.Information)
            msgBox.setText("The Process may take a little while depending on the audio length")
            msgBox.setWindowTitle("NOTIFICATION")
            msgBox.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
            returnValue = msgBox.exec()
            if returnValue == QMessageBox.Ok:
                UIS_RNN.transcribeAudio(self)
        else:
            UIS_RNN.transcribeOnly(self)
        

       
        # self.thread = SpeechThread(self)
        # self.thread.start()
        # Close the loading screen

    def startLive(self):
        self.ui.outputText_Live.setText(' ')
        self.ui.consoleLog.append("> Starting Live  Transcription") 
        UIS_RNN.transcribeLiveStart(self)

    def stopRecording(self):
        UIS_RNN.stopLiveTranscription(self)  # Set stop_recording flag variable to True when you want to stop the recording
        self.ui.btnStartLive.setEnabled(True)
        self.ui.btnStartLive.setStyleSheet("background-color: rgb(12, 149, 53);")
        self.ui.btnStopLive.setEnabled(False)
        self.ui.btnStopLive.setStyleSheet("background-color: #834545;")

    def clearText(self):
        print('clear')
        self.ui.outputText.setText(' ')
    
    def clearTextLive(self):
        print('clear')
        self.ui.outputText_Live.setText(' ')
    
    def exportToDocs(self):
        output = self.ui.outputText.toPlainText()
        doc = docx.Document()
        doc.add_paragraph(output)
        file_dialog = QtWidgets.QFileDialog()
        file_dialog.setAcceptMode(QtWidgets.QFileDialog.AcceptSave)
        file_dialog.setDefaultSuffix("docx")
        file_dialog.setNameFilter("Word Document (*.docx)")

        # Show the Save File Dialog
        if file_dialog.exec_() == QtWidgets.QFileDialog.Accepted:
            file_path = file_dialog.selectedFiles()[0]
            doc.save(file_path)


    def exportPDF(self):
        output = self.ui.outputText.toPlainText()

        # Split the text into lines
        lines = output.split('\n')

        # Create a Save File Dialog
        file_dialog = QtWidgets.QFileDialog()
        file_dialog.setAcceptMode(QtWidgets.QFileDialog.AcceptSave)
        file_dialog.setDefaultSuffix("pdf")
        file_dialog.setNameFilter("PDF Document (*.pdf)")

        # Show the Save File Dialog
        if file_dialog.exec_() == QtWidgets.QFileDialog.Accepted:
            file_path = file_dialog.selectedFiles()[0]

            # Create a PDF document
            pdf = FPDF()
            pdf.add_page()
            pdf.set_left_margin(20)  # Set the left margin (in points)
            pdf.set_font("Arial", size=12)

            # Add each line of text to the PDF
            for line in lines:
                pdf.multi_cell(170, 10, txt=line, align="L")
                pdf.ln(2)  # Add a small space between lines

            pdf.output(file_path, "F")

    def exportToDocsLive(self):
        output = self.ui.outputText_Live.toPlainText()
        doc = docx.Document()
        doc.add_paragraph(output)
        file_dialog = QtWidgets.QFileDialog()
        file_dialog.setAcceptMode(QtWidgets.QFileDialog.AcceptSave)
        file_dialog.setDefaultSuffix("docx")
        file_dialog.setNameFilter("Word Document (*.docx)")

        # Show the Save File Dialog
        if file_dialog.exec_() == QtWidgets.QFileDialog.Accepted:
            file_path = file_dialog.selectedFiles()[0]
            doc.save(file_path)


    def exportPDFLive(self):
        output = self.ui.outputText_Live.toPlainText()

        # Split the text into lines
        lines = output.split('\n')

        # Create a Save File Dialog
        file_dialog = QtWidgets.QFileDialog()
        file_dialog.setAcceptMode(QtWidgets.QFileDialog.AcceptSave)
        file_dialog.setDefaultSuffix("pdf")
        file_dialog.setNameFilter("PDF Document (*.pdf)")

        # Show the Save File Dialog
        if file_dialog.exec_() == QtWidgets.QFileDialog.Accepted:
            file_path = file_dialog.selectedFiles()[0]

            # Create a PDF document
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            # Add each line of text to the PDF
            for line in lines:
                pdf.cell(190, 10, txt=line, ln=1, align="L")

            pdf.output(file_path, "F")




    def exportToSRT(self):
        output = self.ui.outputText.toPlainText()

        if self.ui.enableSpeakerDiarization.isChecked():
            srt = self.transcription_to_srt(output)
            print('SPEAKER DIARIZATION IS ON')
        else:
            srt = self.transcription_to_srt_nospeaker(output)
            print('SPEAKER DIARIZATION IS OFF')

        file_dialog = QtWidgets.QFileDialog()
        file_dialog.setAcceptMode(QtWidgets.QFileDialog.AcceptSave)
        file_dialog.setDefaultSuffix("srt")
        file_dialog.setNameFilter("SRT Subtitle (*.srt)")

        # Show the Save File Dialog
        if file_dialog.exec_() == QtWidgets.QFileDialog.Accepted:
            file_path = file_dialog.selectedFiles()[0]
            with open(file_path, 'w') as f:
                f.write(srt)


    def transcription_to_srt_nospeaker(self,transcription):
        lines = transcription.split('\n')
        srt = ''
        count = 1
        i = 0

        while i < len(lines):
            time_line = lines[i].strip()
            if time_line.startswith('(') and time_line.endswith(')'):
                times = time_line[1:-1]
                if ' - ' in times:
                    start_time, end_time = times.split(' - ')

                    # Corrected regex pattern to match time format without milliseconds
                    start_time = re.findall(r'\d{1,2}:\d{2}:\d{2}', start_time)[0]
                    end_time = re.findall(r'\d{1,2}:\d{2}:\d{2}', end_time)[0]

                    start_h, start_m, start_s = start_time.split(':')
                    start_time_secs = int(start_h) * 3600 + int(start_m) * 60 + float(start_s)
                    end_h, end_m, end_s = end_time.split(':')
                    end_time_secs = int(end_h) * 3600 + int(end_m) * 60 + float(end_s)

                    i += 1
                    # Skip empty lines
                    while i < len(lines) and lines[i].strip() == '':
                        i += 1

                    if i < len(lines):
                        transcription_line = lines[i].strip()
                        subtitle_lines = self.split_into_subtitle_lines(transcription_line)

                        # Calculate the duration per subtitle line
                        duration_per_line = (end_time_secs - start_time_secs) / len(subtitle_lines)

                        # Add SRT entries for each subtitle line
                        for subtitle_index in range(len(subtitle_lines)):
                            start_time_srt = self.format_time(start_time_secs + subtitle_index * duration_per_line)
                            end_time_srt = self.format_time(start_time_secs + (subtitle_index + 1) * duration_per_line)
                            subtitle_line = subtitle_lines[subtitle_index]
                            srt += '{}\n{} --> {}\n{}\n\n'.format(count, start_time_srt, end_time_srt, subtitle_line)
                            count += 1
                i += 1
            else:
                i += 1

        return srt






    

    def transcription_to_srt(self, transcription):
        lines = transcription.split('\n')
        srt = ''
        count = 1
        speakers = {}

        # Find the first and last speakers
        first_speaker = lines[0][:lines[0].find('(')].strip()
        last_speaker = lines[-2][:lines[-2].find('(')].strip()

        # Loop through all lines and add them to the SRT if they contain a speaker name
        for i in range(len(lines)):
            speaker_line = lines[i].strip()
            if speaker_line.startswith('Speaker'):
                speaker_name = speaker_line[speaker_line.find(' ')+1:speaker_line.find(':')]
                speaker_times = speaker_line[speaker_line.find('(')+1:speaker_line.find(')')]
                if ' - ' in speaker_times:
                    start_time, end_time = speaker_times.split(' - ')
                    start_h, start_m, start_s = start_time.split(':')
                    start_time_secs = int(start_h) * 3600 + int(start_m) * 60 + float(start_s)
                    end_h, end_m, end_s = end_time.split(':')
                    end_time_secs = int(end_h) * 3600 + int(end_m) * 60 + float(end_s)

                    # Extract transcription line and split into subtitle lines
                    transcription_line = lines[i+1].strip()
                    subtitle_lines = self.split_into_subtitle_lines(transcription_line)

                    # Calculate the duration per subtitle line
                    duration_per_line = (end_time_secs - start_time_secs) / len(subtitle_lines)
                    if len(subtitle_lines) > 1:
                        end_time_secs = start_time_secs + len(subtitle_lines) * duration_per_line

                    # Add speaker to dictionary if not already there
                    if speaker_name not in speakers:
                        speakers[speaker_name] = len(speakers) + 1

                    # Add SRT entries for each subtitle line
                    for subtitle_index in range(len(subtitle_lines)):
                        start_time_srt = self.format_time(start_time_secs + subtitle_index * duration_per_line)
                        end_time_srt = self.format_time(start_time_secs + (subtitle_index + 1) * duration_per_line)
                        subtitle_line = subtitle_lines[subtitle_index]
                        srt += '{}\n{} --> {}\n{}\n\n'.format(count, start_time_srt, end_time_srt, subtitle_line)
                        count += 1

        # Add empty entries for any missing speakers
        for speaker_name in speakers:
            if speaker_name != first_speaker and speaker_name != last_speaker and 'Speaker {}'.format(speakers[speaker_name]) not in srt:
                srt += '{}\n{} --> {}\n\n\n'.format(count, '00:00:00,000', '00:00:00,000',  '')
                count += 1

        return srt
    
    def split_into_subtitle_lines(self, transcription_line):
        # Split a transcription line into subtitle lines, each with at most 8 words
        subtitle_lines = []
        words = transcription_line.split()
        while len(words) > 0:
            subtitle_line = ''
            while len(words) > 0 and len(subtitle_line.split()) < 10:
                subtitle_line += words.pop(0) + ' '
            subtitle_lines.append(subtitle_line.strip())
        return subtitle_lines

    def format_time(self, time_secs):
        # Format time in SRT format
        h = int(time_secs // 3600)
        m = int((time_secs % 3600) // 60)
        s = time_secs % 60
        ms = int((s - int(s)) * 1000)
        s = int(s)
        return '{:02d}:{:02d}:{:02d},{:03d}'.format(h, m, s, ms)



if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    sys.exit(app.exec_())
